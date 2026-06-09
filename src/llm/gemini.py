"""Gemini async LLM wrapper — replaces claude.py.

Public API (same interface as claude.py):
    embed_text(text)                          -> list[float]  (768-dim)
    score_job(resume, jd)                     -> ScoringResult
    tailor_resume(resume, jd)                 -> str
    write_cover_letter(resume, jd, co, title) -> str
    parse_resume_file(bytes, mime_type)       -> str  (PDF → text via Gemini)

Models used:
    Generation : GEMINI_MODEL env var (default gemini-2.0-flash)
    Embeddings : text-embedding-004  (768-dim)
    PDF parse  : gemini-2.0-flash (multimodal)
"""
from __future__ import annotations

import asyncio
import hashlib
import json
import logging
from dataclasses import dataclass, field

from google import genai
from google.genai import types

from src.config import settings

log = logging.getLogger(__name__)

_client: genai.Client | None = None
_groq_client = None       # primary key client
_groq_client_2 = None     # secondary key client

# In-process embedding cache: sha256(text[:8000]) → list[float]
# Avoids re-embedding the same resume text on every WhatsApp/Gmail request.
# Capped at 256 entries; oldest entry evicted when full.
_EMBED_CACHE: dict[str, list[float]] = {}
_EMBED_CACHE_MAX = 256

# Chunk-level embedding cache: sha256(resume_text[:8000]) → [(chunk_text, embedding)]
# Pre-computed once per unique resume; all retrieval calls reuse this.
_CHUNK_CACHE: dict[str, list[tuple[str, list[float]]]] = {}


def _get_client() -> genai.Client:
    global _client
    if _client is None:
        _client = genai.Client(api_key=settings.gemini_api_key)
    return _client


def _get_groq_client(secondary: bool = False):
    global _groq_client, _groq_client_2
    from groq import AsyncGroq
    if secondary:
        if not settings.groq_api_key_2:
            raise RuntimeError("GROQ_API_KEY_2 not set — no secondary Groq key available")
        if _groq_client_2 is None:
            _groq_client_2 = AsyncGroq(api_key=settings.groq_api_key_2)
        return _groq_client_2
    if not settings.groq_configured:
        raise RuntimeError("GROQ_API_KEY not set — cannot use Groq fallback")
    if _groq_client is None:
        _groq_client = AsyncGroq(api_key=settings.groq_api_key)
    return _groq_client


def _is_groq_rate_limit(exc: Exception) -> bool:
    msg = str(exc)
    return "429" in msg or "rate_limit_exceeded" in msg or "Rate limit" in msg


def _friendly_rate_limit_message() -> str:
    return (
        "⏳ *AI service is temporarily busy*\n\n"
        "Both AI providers have hit their rate limits right now. "
        "This usually clears in 30–60 minutes.\n\n"
        "Your job details have been saved — just send the message again shortly "
        "and the draft will be generated."
    )


async def _groq_generate(system: str, user: str) -> str:
    """Call Groq — tries primary key first, falls back to secondary on 429."""
    for use_secondary in (False, True):
        if use_secondary and not settings.groq_api_key_2:
            break
        try:
            client = _get_groq_client(secondary=use_secondary)
            resp = await client.chat.completions.create(
                model=settings.groq_model,
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                max_tokens=4096,
                temperature=0.3,
            )
            return resp.choices[0].message.content
        except Exception as exc:
            if _is_groq_rate_limit(exc):
                if use_secondary:
                    log.warning("Both Groq keys rate-limited")
                    raise RuntimeError("rate_limit_both_keys") from exc
                log.warning("Primary Groq key rate-limited — trying secondary key")
                continue
            raise
    raise RuntimeError("rate_limit_both_keys")


def _is_gemini_retriable(exc: Exception) -> bool:
    msg = str(exc)
    return "429" in msg or "503" in msg or "500" in msg or "RESOURCE_EXHAUSTED" in msg or "UNAVAILABLE" in msg


# ---------------------------------------------------------------------------
# Shared dataclass (same as claude.py so agents need no changes)
# ---------------------------------------------------------------------------

@dataclass
class ScoringResult:
    score: float
    reasoning: str
    decision: str
    strengths: list[str] = field(default_factory=list)
    gaps: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Embeddings  (768-dim, text-embedding-004)
# ---------------------------------------------------------------------------

async def embed_text(text: str) -> list[float]:
    """Embed text to a 768-dim vector using gemini-embedding-001. Results are cached in-process."""
    snippet = text[:8000]
    key = hashlib.sha256(snippet.encode()).hexdigest()
    if key in _EMBED_CACHE:
        return _EMBED_CACHE[key]

    client = _get_client()
    result = await client.aio.models.embed_content(
        model="models/gemini-embedding-001",
        contents=snippet,
        config={"output_dimensionality": 768},
    )
    vec = list(result.embeddings[0].values)

    if len(_EMBED_CACHE) >= _EMBED_CACHE_MAX:
        # Evict oldest entry
        _EMBED_CACHE.pop(next(iter(_EMBED_CACHE)))
    _EMBED_CACHE[key] = vec
    return vec


async def precompute_resume_chunks(resume_text: str) -> list[tuple[str, list[float]]]:
    """
    Split resume into semantic sections and embed each one concurrently.

    Results are keyed by sha256(resume_text[:8000]) so the same resume is never
    embedded twice — not within a process lifetime, and not across calls in the
    same session.  Callers (matching, tailoring, WhatsApp) should call this once
    per resume and let retrieval functions reuse the cached result.

    Returns list of (chunk_text, embedding) in original document order.
    """
    import re as _re

    key = hashlib.sha256(resume_text[:8_000].encode()).hexdigest()
    if key in _CHUNK_CACHE:
        return _CHUNK_CACHE[key]

    raw = _re.split(r"\n{2,}|(?=\n[A-Z][A-Z\s]{3,}:?\n)", resume_text)
    chunks = [c.strip() for c in raw if len(c.strip()) > 40]

    if not chunks:
        result = [(resume_text[:5_000], [])]
        _CHUNK_CACHE[key] = result
        return result

    try:
        embs = await asyncio.gather(*[embed_text(c[:1_000]) for c in chunks])
    except Exception as exc:
        log.warning("Chunk embedding failed, falling back to raw resume: %s", exc)
        result = [(resume_text[:5_000], [])]
        _CHUNK_CACHE[key] = result
        return result

    result = list(zip(chunks, embs))
    _CHUNK_CACHE[key] = result
    return result


def _cosine_similarity(a: list[float], b: list[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    mag_a = sum(x * x for x in a) ** 0.5
    mag_b = sum(x * x for x in b) ** 0.5
    return dot / (mag_a * mag_b) if mag_a and mag_b else 0.0


async def get_relevant_resume_context(
    resume_text: str,
    job_description: str,
    top_k: int = 6,
    max_chars: int = 5_000,
) -> str:
    """
    Return the top_k resume sections most relevant to job_description.

    Chunk embeddings are computed via precompute_resume_chunks() and cached
    in _CHUNK_CACHE — so for N jobs against the same resume only one set of
    embedding API calls is made (first call), all subsequent calls are free.
    """
    cached = await precompute_resume_chunks(resume_text)

    # Bail out if we got the fallback (no real embeddings)
    if not cached or not cached[0][1]:
        return resume_text[:max_chars]

    try:
        jd_emb = await embed_text(job_description[:2_000])
    except Exception as exc:
        log.warning("JD embedding failed, returning cached first chunks: %s", exc)
        return "\n\n".join(c for c, _ in cached[:top_k])[:max_chars]

    chunks = [c for c, _ in cached]
    embs   = [e for _, e in cached]

    scored = sorted(
        zip(chunks, embs),
        key=lambda pair: _cosine_similarity(jd_emb, pair[1]),
        reverse=True,
    )

    # Take top_k chunks; preserve original document order for readability
    top_set = {c for c, _ in scored[:top_k]}
    ordered = [c for c in chunks if c in top_set]

    result = "\n\n".join(ordered)
    return result[:max_chars]


async def classify_email_as_job_opportunity(subject: str, body: str) -> bool:
    """Backward-compatible shim — returns True if the email is a hiring/interview email."""
    is_job, _ = await classify_and_type_email(subject, body)
    return is_job


async def classify_and_type_email(
    subject: str, body: str, sender: str = ""
) -> tuple[bool, str]:
    """
    Classify an email in a single AI call.

    Returns:
        (needs_draft, email_type)
        needs_draft — True only if a reply draft should be created
        email_type  — "interview" | "opportunity" | "rejection" | "other"

    Only "interview" and "opportunity" produce drafts.
    "rejection" and "other" are silently skipped.
    """
    prompt = f"""You are classifying an email to decide if a reply draft needs to be created.

SUBJECT: {subject}
FROM: {sender}
BODY (first 1200 chars):
{body[:1200]}

Classify as exactly one of these four types:

INTERVIEW — An email from a real person inviting or scheduling the recipient for a job interview, assessment, or technical test for a specific role.

OPPORTUNITY — A recruiter or hiring manager personally reaching out about a specific job role, asking if the recipient is interested. NOT a mass blast or digest.

REJECTION — A polite rejection, regret, or "we went with another candidate" email. No reply is needed.

OTHER — Everything else: job alert digests, newsletters, automated ATS confirmations ("We received your application"), application status updates with no next step, promotional emails, LinkedIn notifications, mass marketing, general company updates.

Rules — classify as OTHER if ANY of these apply:
- Sender address contains: noreply, no-reply, donotreply, alerts, notifications, mailer
- Subject starts with: "Job Alert", "New jobs for you", "Jobs matching", "Your application", "Application received"
- No specific role or company is personally mentioned
- Reads like a template blast with no personal greeting
- No actionable next step for the recipient

Rules — classify as REJECTION if:
- Contains phrases like "we have decided", "we will not be moving forward", "not selected", "went with another candidate", "position has been filled", "unfortunately", "regret to inform"

Reply with EXACTLY one word: INTERVIEW, OPPORTUNITY, REJECTION, or OTHER"""

    raw: str | None = None
    try:
        client = _get_client()
        if client:
            resp = await client.aio.models.generate_content(
                model=settings.gemini_model,
                contents=prompt,
                config=types.GenerateContentConfig(max_output_tokens=10, temperature=0.0),
            )
            raw = (resp.text or "").strip().upper()
    except Exception as exc:
        log.warning("Email classification (Gemini) failed: %s — trying Groq", exc)

    if not raw and settings.groq_configured:
        try:
            raw = (await _groq_generate(
                "You classify emails. Reply with exactly one word: INTERVIEW, OPPORTUNITY, REJECTION, or OTHER.",
                prompt,
            ) or "").strip().upper()
        except Exception as exc:
            log.warning("Email classification (Groq) also failed: %s — defaulting to OTHER", exc)

    if raw and raw.startswith("INTERVIEW"):
        return True, "interview"
    if raw and raw.startswith("OPPORTUNITY"):
        return True, "opportunity"
    if raw and raw.startswith("REJECTION"):
        return False, "rejection"
    # Default fail-safe: don't create drafts for unclassifiable mail
    return False, "other"


async def write_interview_reply(
    profile: dict,
    email_subject: str,
    email_body: str,
    job: dict,
) -> str:
    """
    Write a personalized interview confirmation reply.
    `profile["resume_text"]` should already be embedding-ranked for this role
    so the model sees the most relevant experience, not a raw first-N-chars slice.
    """
    name      = profile.get("full_name", "")
    email_addr = profile.get("email", "")
    phone     = profile.get("phone", "")
    resume    = (profile.get("resume_text") or "")[:3_000]
    role      = job.get("job_title") or "the position"
    company   = job.get("company") or "your company"
    hr_name   = job.get("hr_name") or "Hiring Manager"

    system = (
        "You write professional, concise email replies on behalf of a job candidate. "
        "Never invent facts. Base every claim on the resume provided. "
        "Sound warm and confident — never generic or robotic."
    )
    prompt = f"""Write a personalized interview confirmation reply for the candidate below.

CANDIDATE:
Name: {name}
{f"Email: {email_addr}" if email_addr else ""}
{f"Phone: {phone}" if phone else ""}

MOST RELEVANT RESUME SECTIONS (selected by semantic similarity to this role):
{resume}

ORIGINAL EMAIL:
Subject: {email_subject}
{email_body[:1_500]}

ROLE: {role} at {company}
ADDRESSED TO: {hr_name}

Write a warm, professional reply of 3 short paragraphs:

Paragraph 1 — Thank them for the interview invitation. Express genuine enthusiasm for the {role} role at {company}. Be specific about what excites you (drawn from the job description in the original email).

Paragraph 2 — In 2-3 sentences, mention 1-2 specific experiences from the resume that are directly relevant to this role. Use concrete details (project names, technologies, outcomes) — not generic skill lists. This shows you are a strong fit before the interview even starts.

Paragraph 3 — Confirm you are happy to proceed. Ask for the preferred date, time, and format (in-person / video / phone). Offer your availability or ask them to share a scheduling link. Keep it brief and action-oriented.

Total target: 150–200 words. No greeting line or sign-off — return only the body paragraphs."""

    text: str | None = None
    try:
        client = _get_client()
        if client:
            resp = await client.aio.models.generate_content(
                model=settings.gemini_model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    system_instruction=system,
                    max_output_tokens=512,
                    temperature=0.6,
                ),
            )
            text = (resp.text or "").strip()
    except Exception as exc:
        log.warning("write_interview_reply (Gemini) failed: %s — trying Groq", exc)

    if not text and settings.groq_configured:
        try:
            text = (await _groq_generate(system, prompt) or "").strip()
        except Exception as exc:
            log.warning("write_interview_reply (Groq) also failed: %s", exc)

    return text or (
        f"Thank you for inviting me to interview for the {role} role at {company}. "
        "I am very enthusiastic about this opportunity and would be happy to proceed.\n\n"
        "Please let me know your preferred date, time, and interview format "
        "(in-person / video call / phone) and I will confirm my availability."
    )


# ---------------------------------------------------------------------------
# Job scoring
# ---------------------------------------------------------------------------

_SCORE_SYSTEM = """\
You are a professional recruiter. Given a candidate resume and a job description,
evaluate fit on a 0–10 scale. Return ONLY valid JSON matching this schema:
{
  "score_0_to_10": <integer 0-10>,
  "reasoning": "<2-3 sentence summary>",
  "decision": "apply" | "skip",
  "strengths": ["...", ...],
  "gaps": ["...", ...]
}
Decision rule: score >= 6 → apply, else skip.
NEVER fabricate experience the candidate does not have.
"""


async def score_job(resume_text: str, job_description: str) -> ScoringResult:
    # Select the most relevant resume sections for this JD (cheap after first call —
    # chunk embeddings are cached so each subsequent job only costs 1 JD embedding).
    try:
        resume_ctx = await get_relevant_resume_context(
            resume_text, job_description, top_k=4, max_chars=2_500,
        )
    except Exception:
        resume_ctx = resume_text[:2_500]

    prompt = (
        f"<resume>\n{resume_ctx}\n</resume>\n\n"
        f"<job_description>\n{job_description[:3_000]}\n</job_description>"
    )
    text: str | None = None
    try:
        client = _get_client()
        response = await client.aio.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
            config=types.GenerateContentConfig(system_instruction=_SCORE_SYSTEM),
        )
        text = response.text
    except Exception as exc:
        if settings.groq_configured and _is_gemini_retriable(exc):
            log.warning("Gemini score_job failed (%s) — falling back to Groq", exc)
            text = await _groq_generate(_SCORE_SYSTEM, prompt)
        else:
            raise
    if not text or not text.strip():
        log.warning("score_job received empty response — returning fallback")
        return ScoringResult(score=0.0, reasoning="empty response", decision="skip")
    data = json.loads(_extract_json(text))
    raw = int(data.get("score_0_to_10", 0))
    return ScoringResult(
        score=raw * 10,  # store as 0-100
        reasoning=data.get("reasoning", ""),
        decision=data.get("decision", "skip"),
        strengths=data.get("strengths", []),
        gaps=data.get("gaps", []),
    )


# ---------------------------------------------------------------------------
# Resume tailoring
# ---------------------------------------------------------------------------

_TAILOR_SYSTEM = """\
You are an expert resume writer.
Tailor the given resume so it best matches the job description.
Rules:
- NEVER invent skills, experience, or qualifications the candidate does not have.
- Only reword, reorder, and emphasise existing content.
- Keep the same format (plain text, sections with headings).
- Output ONLY the tailored resume text — no commentary.
"""


async def tailor_resume(resume_text: str, job_description: str) -> str:
    # For tailoring we want more chunks (high top_k) so nothing important is pruned,
    # but still skip sections that are irrelevant to this specific JD.
    if len(resume_text) > 3_000:
        try:
            resume_ctx = await get_relevant_resume_context(
                resume_text, job_description, top_k=10, max_chars=6_000,
            )
        except Exception:
            resume_ctx = resume_text[:6_000]
    else:
        resume_ctx = resume_text

    prompt = (
        f"<resume>\n{resume_ctx}\n</resume>\n\n"
        f"<job_description>\n{job_description[:4_000]}\n</job_description>"
    )
    text: str | None = None
    try:
        client = _get_client()
        response = await client.aio.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=_TAILOR_SYSTEM,
                max_output_tokens=4096,
            ),
        )
        text = response.text
    except Exception as exc:
        if settings.groq_configured and _is_gemini_retriable(exc):
            log.warning("Gemini tailor_resume failed (%s) — falling back to Groq", exc)
            text = await _groq_generate(_TAILOR_SYSTEM, prompt)
        else:
            raise
    if not text or not text.strip():
        log.warning("tailor_resume received empty response — returning original")
        return resume_text
    return text


# ---------------------------------------------------------------------------
# Cover letter
# ---------------------------------------------------------------------------

_COVER_SYSTEM = """\
You are an expert job application writer.
Write a concise, compelling cover letter (3–4 paragraphs, ≤ 400 words).
Rules:
- Address it to the hiring team at the company.
- Reference specific requirements from the JD.
- NEVER fabricate experience or qualifications.
- Output ONLY the cover letter text — no subject line, no commentary.
"""


async def write_cover_letter(
    resume_text: str,
    job_description: str,
    company: str,
    job_title: str,
) -> str:
    try:
        resume_ctx = await get_relevant_resume_context(
            resume_text, job_description, top_k=5, max_chars=3_000,
        )
    except Exception:
        resume_ctx = resume_text[:3_000]

    prompt = (
        f"Company: {company}\nRole: {job_title}\n\n"
        f"<resume>\n{resume_ctx}\n</resume>\n\n"
        f"<job_description>\n{job_description[:3_000]}\n</job_description>"
    )
    try:
        client = _get_client()
        response = await client.aio.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
            config=types.GenerateContentConfig(system_instruction=_COVER_SYSTEM),
        )
        return response.text
    except Exception as exc:
        if settings.groq_configured and _is_gemini_retriable(exc):
            log.warning("Gemini write_cover_letter failed (%s) — falling back to Groq", exc)
            return await _groq_generate(_COVER_SYSTEM, prompt)
        raise


# ---------------------------------------------------------------------------
# Application email  (richer than cover letter — used by WhatsApp + email-drafts)
# ---------------------------------------------------------------------------

_APP_EMAIL_SYSTEM = """\
You are a senior professional job application email writer used by real candidates applying for real jobs.

Rules:
- Use ONLY facts from the resume — never fabricate skills, titles, or achievements.
- Be specific: name the company and role, reference concrete details from the JD.
- Tone: confident, direct, human — like a strong candidate who knows their value.
- Length: 200-280 words across exactly 4 short, punchy paragraphs.
- Each paragraph: 2-3 tight sentences. No padding, no filler.
- Output ONLY the four body paragraphs — no greeting line, no sign-off, no subject line.
- Forbidden openers: "I am writing to apply", "I am excited to", "I believe I would be", "I am pleased to".
- Forbidden closers: "Thank you for considering my application", "I look forward to hearing from you at your earliest convenience".
- Use full sentences, no bullet points inside the email body.
- If the resume contains metrics (years, percentages, team sizes, project outcomes) — use them.
"""


async def write_application_email(profile: dict, job: dict) -> str:
    """Write a professional, full-length application email body."""
    name         = profile.get("full_name", "")
    email_addr   = profile.get("email", "")
    phone        = profile.get("phone", "")
    target_title = profile.get("target_title", "")
    resume       = (profile.get("resume_text") or "")[:8_000]

    company   = job.get("company") or "the company"
    job_title = job.get("job_title") or "the role"
    location  = job.get("location") or job.get("target_location") or ""
    jd_text   = (job.get("description") or "")[:3_000]
    hr_name   = job.get("hr_name") or "Hiring Team"

    prompt = f"""Write a professional job application email body for the applicant below.
This will be sent directly to the hiring manager — it must be polished, specific, and compelling.

APPLICANT:
Name: {name}
Email: {email_addr}
Phone: {phone}
{f"Target Role: {target_title}" if target_title else ""}

FULL RESUME:
{resume}

ROLE BEING APPLIED FOR:
Company: {company}
Position: {job_title}
{f"Location: {location}" if location else ""}
Job Description:
{jd_text}

Write exactly 4 short paragraphs addressed to {hr_name}. Each paragraph is 2-3 sentences maximum:

Paragraph 1 — Opening (2 sentences): State the exact role. Make one sharp, specific observation about {company} or the role that shows you've read the JD — not a generic compliment.

Paragraph 2 — Core value (3 sentences): Lead with your single strongest relevant achievement from the resume. Use a specific number or outcome if available. Connect it directly to what the JD is asking for.

Paragraph 3 — Skills fit (2-3 sentences): Name 2 specific skills or experiences from the resume that match requirements in the JD. Be concrete, not vague.

Paragraph 4 — Close (2 sentences): Say you've attached your CV and invite a conversation. Be direct — no "I look forward to hearing from you at your earliest convenience".

Target 200-250 words total. Every sentence earns its place. Sound human and confident — not like a template."""

    text: str | None = None
    try:
        client = _get_client()
        response = await client.aio.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=_APP_EMAIL_SYSTEM,
                max_output_tokens=1024,
                temperature=0.7,
            ),
        )
        text = response.text
    except Exception as exc:
        if settings.groq_configured and _is_gemini_retriable(exc):
            log.warning("Gemini write_application_email failed (%s) — falling back to Groq", exc)
            text = await _groq_generate(_APP_EMAIL_SYSTEM, prompt)
        else:
            raise
    if not text or not text.strip():
        raise ValueError("write_application_email received an empty response from all AI providers")
    return text.strip()


# ---------------------------------------------------------------------------
# Resume file parsing  (PDF → plain text via Gemini multimodal)
# ---------------------------------------------------------------------------

async def parse_resume_file(file_bytes: bytes, mime_type: str, filename: str = "") -> str:
    """Extract structured plain text from an uploaded resume file (PDF/DOCX/DOC/TXT)."""
    client = _get_client()
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""

    # ── PDF: Gemini multimodal, pypdf fallback ─────────────────
    if mime_type == "application/pdf" or ext == ".pdf":
        try:
            part = types.Part.from_bytes(data=file_bytes, mime_type="application/pdf")
            response = await client.aio.models.generate_content(
                model=settings.gemini_model,
                contents=[
                    part,
                    "Extract all text from this resume/CV. Return only the plain text, "
                    "preserving sections and headings. Do not add commentary.",
                ],
            )
            text = (response.text or "").strip()
            if text:
                return text
        except Exception as exc:
            log.warning("Gemini PDF parse failed, trying pypdf: %s", exc)

        # pypdf fallback — clean text, no null bytes
        try:
            import pypdf
            from io import BytesIO
            reader = pypdf.PdfReader(BytesIO(file_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages).strip()
            if text:
                return text
        except Exception as exc:
            log.warning("pypdf fallback failed: %s", exc)

        return ""  # let upload_resume handle the empty-text case

    # ── DOCX: python-docx paragraph extraction ──────────────────
    if ext in (".docx",) or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(file_bytes))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        lines.append(row_text)
            text = "\n".join(lines).strip()
            if text:
                return text
        except Exception:
            pass  # fall through to raw decode

    # ── DOC (legacy binary): python-docx attempt then raw ───────
    if ext == ".doc" or mime_type == "application/msword":
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(file_bytes))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(lines).strip()
            if text:
                return text
        except Exception:
            pass  # old binary .doc — fall through to raw bytes

    # ── Plain text / fallback ────────────────────────────────────
    try:
        return file_bytes.decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------

def _extract_json(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        inner = [l for l in lines[1:] if l.strip() != "```"]
        return "\n".join(inner).strip()
    start = text.find("{")
    end = text.rfind("}") + 1
    if start >= 0 and end > start:
        return text[start:end]
    return text
